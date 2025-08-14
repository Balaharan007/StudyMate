import os
import requests
import PyPDF2
import docx
from io import BytesIO
from typing import List, Dict, Any, Union
import google.generativeai as genai
from dotenv import load_dotenv
import hashlib
import json
import re
import numpy as np

load_dotenv()

class SimpleDocumentProcessor:
    def __init__(self):
        # Initialize Gemini
        genai.configure(api_key=os.getenv("GEMINI_API_KEY"))
        self.gemini_model = genai.GenerativeModel('gemini-2.0-flash-exp')
        
        # Simple in-memory storage for processed documents
        self.document_chunks = {}
    
    def download_document(self, url: str) -> bytes:
        """Download document from URL"""
        try:
            response = requests.get(url, timeout=30)
            response.raise_for_status()
            return response.content
        except Exception as e:
            raise Exception(f"Failed to download document: {str(e)}")
    
    def extract_text_from_pdf(self, content: Union[bytes, str]) -> str:
        """Extract text from PDF"""
        try:
            if isinstance(content, str):
                with open(content, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text() + "\n"
            else:
                pdf_file = BytesIO(content)
                reader = PyPDF2.PdfReader(pdf_file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
            return text
        except Exception as e:
            raise Exception(f"Failed to extract PDF text: {str(e)}")
    
    def extract_text_from_docx(self, content: Union[bytes, str]) -> str:
        """Extract text from DOCX"""
        try:
            if isinstance(content, str):
                doc = docx.Document(content)
            else:
                doc = docx.Document(BytesIO(content))
            
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            raise Exception(f"Failed to extract DOCX text: {str(e)}")
    
    def extract_text_from_txt(self, content: Union[bytes, str]) -> str:
        """Extract text from TXT file"""
        try:
            if isinstance(content, str):
                with open(content, 'r', encoding='utf-8') as file:
                    return file.read()
            else:
                return content.decode('utf-8')
        except Exception as e:
            raise Exception(f"Failed to extract TXT text: {str(e)}")
    
    def process_uploaded_file(self, file_content: bytes, filename: str) -> str:
        """Process uploaded file content"""
        file_ext = filename.lower().split('.')[-1]
        
        if file_ext == 'pdf':
            return self.extract_text_from_pdf(file_content)
        elif file_ext in ['docx', 'doc']:
            return self.extract_text_from_docx(file_content)
        elif file_ext == 'txt':
            return self.extract_text_from_txt(file_content)
        else:
            try:
                return file_content.decode('utf-8')
            except:
                raise Exception(f"Unsupported file format: {file_ext}")
    
    def extract_text(self, source: str, is_file_path: bool = False) -> str:
        """Extract text from document URL or file path"""
        print(f"Processing source: {source}")
        print(f"Is file path: {is_file_path}")
        
        if is_file_path:
            if source.lower().endswith('.pdf'):
                return self.extract_text_from_pdf(source)
            elif source.lower().endswith(('.docx', '.doc')):
                return self.extract_text_from_docx(source)
            elif source.lower().endswith('.txt'):
                return self.extract_text_from_txt(source)
            else:
                raise Exception("Unsupported file format")
        else:
            try:
                content = self.download_document(source)
                print(f"Downloaded content size: {len(content)} bytes")
                print(f"Content header: {content[:50]}")
                
                # Extract file extension from URL (handle query parameters)
                url_path = source.split('?')[0]  # Remove query parameters
                print(f"URL path without query: {url_path}")
                
                # Check for PDF signature
                is_pdf = b'%PDF' in content[:100] or url_path.lower().endswith('.pdf') or 'pdf' in source.lower()
                print(f"Is PDF: {is_pdf}")
                
                if is_pdf:
                    return self.extract_text_from_pdf(content)
                elif url_path.lower().endswith(('.docx', '.doc')) or 'docx' in source.lower() or 'doc' in source.lower():
                    return self.extract_text_from_docx(content)
                elif url_path.lower().endswith('.txt') or 'txt' in source.lower():
                    return self.extract_text_from_txt(content)
                else:
                    # For Azure blob URLs, try PDF first, then text
                    if 'blob.core.windows.net' in source:
                        try:
                            return self.extract_text_from_pdf(content)
                        except:
                            pass
                    
                    # Try as text
                    try:
                        text_content = content.decode('utf-8')
                        if len(text_content.strip()) > 0:
                            return text_content
                    except:
                        pass
                    
                    raise Exception(f"Unsupported document format. Content type not recognized from URL: {url_path}")
            except Exception as e:
                print(f"Error in extract_text: {str(e)}")
                raise
    
    def chunk_text(self, text: str, chunk_size: int = 1500, overlap: int = 300) -> List[Dict]:
        """Split text into chunks with metadata"""
        text = re.sub(r'\s+', ' ', text).strip()
        
        # Split into sentences
        sentences = re.split(r'[.!?]+', text)
        
        chunks = []
        current_chunk = ""
        chunk_id = 0
        start_pos = 0
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
                
            if len(current_chunk) + len(sentence) > chunk_size and current_chunk:
                chunks.append({
                    "id": f"chunk_{chunk_id}",
                    "text": current_chunk.strip(),
                    "start_pos": start_pos,
                    "end_pos": start_pos + len(current_chunk)
                })
                
                overlap_text = current_chunk[-overlap:] if len(current_chunk) > overlap else current_chunk
                current_chunk = overlap_text + " " + sentence
                start_pos = start_pos + len(current_chunk) - len(overlap_text) - len(sentence)
                chunk_id += 1
            else:
                current_chunk += " " + sentence if current_chunk else sentence
        
        if current_chunk.strip():
            chunks.append({
                "id": f"chunk_{chunk_id}",
                "text": current_chunk.strip(),
                "start_pos": start_pos,
                "end_pos": start_pos + len(current_chunk)
            })
        
        return chunks
    
    def simple_similarity(self, query: str, text: str) -> float:
        """Simple text similarity using word overlap"""
        query_words = set(query.lower().split())
        text_words = set(text.lower().split())
        
        if not query_words or not text_words:
            return 0.0
        
        intersection = query_words.intersection(text_words)
        union = query_words.union(text_words)
        
        return len(intersection) / len(union) if union else 0.0
    
    def process_document(self, source: str, is_file_path: bool = False, file_content: bytes = None, filename: str = None) -> Dict:
        """Complete document processing pipeline"""
        try:
            # Extract text
            if file_content and filename:
                text = self.process_uploaded_file(file_content, filename)
                document_id = hashlib.md5((filename + str(len(file_content))).encode()).hexdigest()
            else:
                text = self.extract_text(source, is_file_path)
                document_id = hashlib.md5(source.encode()).hexdigest()
            
            # Chunk text
            chunks = self.chunk_text(text)
            
            # Store chunks in memory
            self.document_chunks[document_id] = chunks
            
            return {
                "success": True,
                "text": text,
                "chunks": len(chunks),
                "document_id": document_id,
                "message": f"Successfully processed document with {len(chunks)} chunks"
            }
        
        except Exception as e:
            return {
                "success": False,
                "error": str(e)
            }
    
    def search_similar_chunks(self, query: str, document_id: str = None, top_k: int = 5) -> List[Dict]:
        """Search for similar chunks using simple text similarity"""
        try:
            results = []
            
            # Search in specific document or all documents
            if document_id and document_id in self.document_chunks:
                chunks = self.document_chunks[document_id]
                for chunk in chunks:
                    score = self.simple_similarity(query, chunk["text"])
                    if score > 0:
                        results.append({
                            "score": score,
                            "text": chunk["text"],
                            "chunk_id": chunk["id"],
                            "document_id": document_id
                        })
            else:
                # Search in all documents
                for doc_id, chunks in self.document_chunks.items():
                    for chunk in chunks:
                        score = self.simple_similarity(query, chunk["text"])
                        if score > 0:
                            results.append({
                                "score": score,
                                "text": chunk["text"],
                                "chunk_id": chunk["id"],
                                "document_id": doc_id
                            })
            
            # Sort by score and return top_k
            results.sort(key=lambda x: x["score"], reverse=True)
            return results[:top_k]
        
        except Exception as e:
            print(f"Search error: {e}")
            return []
    
    def generate_answer(self, question: str, relevant_chunks: List[Dict]) -> Dict:
        """Generate answer using Gemini"""
        if not relevant_chunks:
            return {
                "answer": "I couldn't find relevant information in the document to answer your question.",
                "relevant_chunks": [],
                "reasoning": "No relevant document sections found"
            }
            
        context = "\n\n".join([f"[Clause {chunk['chunk_id']}]: {chunk['text']}" for chunk in relevant_chunks])
        
        prompt = f"""
        You are an intelligent document analysis agent specializing in insurance, legal, HR, and compliance domains.
        
        Based on the following document clauses and user question, provide a comprehensive and accurate answer.

        DOCUMENT CLAUSES:
        {context}

        USER QUESTION:
        {question}

        INSTRUCTIONS:
        1. Analyze the relevant clauses carefully
        2. Provide a clear, accurate answer based ONLY on the document content
        3. If the document doesn't contain sufficient information, state that clearly
        4. Reference specific clauses when relevant
        5. Be precise and professional
        6. Focus on insurance policies, legal contracts, HR policies, or compliance documents
        7. Provide explainable reasoning for your decision
        8. Answer should be concise but comprehensive
        9. Cite specific clause numbers or sections when possible

        ANSWER:
        """
        
        try:
            response = self.gemini_model.generate_content(prompt)
            return {
                "answer": response.text,
                "relevant_chunks": relevant_chunks,
                "reasoning": f"Answer based on simple text similarity search of {len(relevant_chunks)} document clauses"
            }
        except Exception as e:
            return {
                "answer": f"Error generating answer: {str(e)}",
                "relevant_chunks": relevant_chunks,
                "reasoning": "Error occurred during LLM processing"
            }
